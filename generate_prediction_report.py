"""
Generate Word document: Prediction process steps and model fit evaluation.
"""

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("output")
PREDICTION_DIR = OUTPUT_DIR / "prediction"
REPORT_PATH = OUTPUT_DIR / "Prediction_Process_Report.docx"


def read_csv_safe(path: Path) -> pd.DataFrame | None:
    for enc in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
        try:
            df = pd.read_csv(path, encoding=enc)
            if df.empty or len(df.columns) == 0:
                return None
            return df
        except (UnicodeDecodeError, pd.errors.EmptyDataError, pd.errors.ParserError):
            continue
    return None


def df_to_table(doc: Document, df: pd.DataFrame, style: str = "Table Grid") -> None:
    df = df.fillna("")
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style=style)
    for j, col in enumerate(df.columns):
        table.rows[0].cells[j].text = str(col)
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row in df.iterrows():
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str | None = None, width: float = 5.0) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Image not found: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].italic = True
    doc.add_paragraph()


def main():
    doc = Document()

    doc.add_heading("Dispense Method Prediction Process Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    # --- 1. Steps Taken ---
    doc.add_heading("1. Prediction Process Steps", level=1)

    steps = [
        ("Step 1: Base Data Extraction", "Export fill-level data from OLC (order_lifecycle_nrt) via 01_base_data.sql. Includes dispense_method (target), product attributes, wh_name, order_qty. Date range: prior 6 months (dynamic)."),
        ("Step 2: Feature Selection", "Run feature_selection.py to identify predictive features. Uses chi-square and mutual information to rank features by association with dispense_method. Output: selected_features.txt."),
        ("Step 3: Data Preparation", "Load base_data.csv; encode categorical features (LabelEncoder); prepare train/test split (80/20 stratified by dispense_method)."),
        ("Step 4: Model Training", "Train Random Forest classifier (n_estimators=100, max_depth=15) on training set. Uses selected features; optionally weights by order_qty."),
        ("Step 5: Model Evaluation", "Predict on held-out test set. Compute accuracy, macro F1, weighted F1, confusion matrix, per-class precision/recall."),
        ("Step 6: Outputs", "Save metrics (metrics.json), classification report, confusion matrix (CSV + PNG), feature importance (CSV + PNG)."),
    ]
    for title, desc in steps:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        doc.add_paragraph()

    # --- 2. Model Fit / Testing ---
    doc.add_heading("2. Model Fit Evaluation", level=1)

    metrics_path = PREDICTION_DIR / "metrics.json"
    if metrics_path.exists():
        with open(metrics_path) as f:
            metrics = json.load(f)
        doc.add_heading("Summary Metrics", level=2)
        doc.add_paragraph(f"Accuracy: {metrics['accuracy']:.2%}")
        doc.add_paragraph(f"Macro F1: {metrics['macro_f1']:.4f} (average F1 across classes, handles imbalance)")
        doc.add_paragraph(f"Weighted F1: {metrics['weighted_f1']:.4f} (F1 weighted by class frequency)")
        doc.add_paragraph(f"Training samples: {metrics['n_train']:,}")
        doc.add_paragraph(f"Test samples: {metrics['n_test']:,}")
        doc.add_paragraph(f"Features used: {', '.join(metrics['features'])}")
        doc.add_paragraph(f"Target classes: {', '.join(metrics['classes'])}")
        doc.add_paragraph()
    else:
        doc.add_paragraph("[Run build_prediction_model.py first to generate metrics]")
        doc.add_paragraph()

    doc.add_heading("Classification Report (Per-Class)", level=2)
    report_path = PREDICTION_DIR / "classification_report.csv"
    if report_path.exists():
        df = read_csv_safe(report_path)
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[Report empty or invalid]")
    else:
        doc.add_paragraph("[Classification report not found]")
    doc.add_paragraph()

    doc.add_heading("Confusion Matrix", level=2)
    doc.add_paragraph("Rows = actual dispense_method; Columns = predicted. Diagonal = correct predictions.")
    add_image(doc, PREDICTION_DIR / "confusion_matrix.png", caption="Confusion matrix (test set)", width=5.0)
    cm_path = PREDICTION_DIR / "confusion_matrix.csv"
    if cm_path.exists():
        df = read_csv_safe(cm_path)
        if df is not None:
            df_to_table(doc, df)
    doc.add_paragraph()

    doc.add_heading("Feature Importance", level=2)
    doc.add_paragraph("Random Forest feature importances (Gini importance). Higher = more predictive of dispense_method.")
    add_image(doc, PREDICTION_DIR / "feature_importance.png", caption="Feature importance", width=5.0)
    fi_path = PREDICTION_DIR / "feature_importance.csv"
    if fi_path.exists():
        df = read_csv_safe(fi_path)
        if df is not None:
            df_to_table(doc, df)
    doc.add_paragraph()

    doc.add_heading("3. Next Steps", level=1)
    doc.add_paragraph(
        "• Deploy model for pre-dispense prediction (product_id + wh_name → predicted dispense_method)\n"
        "• For new products not in training data: use product_type + wh_name rules from Part 1 distributions\n"
        "• Consider temporal validation: train on first 5 months, test on 6th month\n"
        "• Monitor accuracy over time; retrain as dispense patterns evolve"
    )

    out_path = REPORT_PATH
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Prediction report saved to {out_path}")


if __name__ == "__main__":
    main()
