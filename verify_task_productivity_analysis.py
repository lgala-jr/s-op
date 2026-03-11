"""
Task Productivity Analysis: histogram distribution of logic_adjusted_new_step_duration_minutes
Broken out by task, wh_name, and other useful cuts.
Generates visuals and a Word document.
"""

import time
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from docx import Document
from docx.shared import Inches
from scipy import stats
from scipy.signal import find_peaks

DATA_PATH = Path("output/verify_task_productivity.csv")
OUTPUT_DIR = Path("output/verify_task_productivity")
DURATION_COL = "duration_minutes"  # Combined: Automated=ceil(ended-started)sec/60, Manual=logic_adjusted
# Subsample for histograms when n > this (faster plotting, shape preserved)
HIST_SAMPLE_MAX = 100_000


def load_data() -> pd.DataFrame:
    for enc in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
        try:
            df = pd.read_csv(DATA_PATH, encoding=enc)
            df.columns = df.columns.str.strip()
            # Normalize Wh_Name -> wh_name for consistency
            if "Wh_Name" in df.columns and "wh_name" not in df.columns:
                df["wh_name"] = df["Wh_Name"]
            # task_mode: Automated if task_performed_by contains 'beyondrx', else Manual
            perf_col = "task_performed_by" if "task_performed_by" in df.columns else "Task_Performed_By"
            if perf_col in df.columns:
                df["task_mode"] = np.where(
                    df[perf_col].astype(str).str.lower().str.contains("beyondrx", na=False),
                    "Automated",
                    "Manual",
                )
            # duration_minutes: Automated = ceil(task_ended - task_started) to 1 sec, then /60; Manual = logic_adjusted
            df["duration_minutes"] = _compute_duration_minutes(df)
            return df
        except (UnicodeDecodeError, FileNotFoundError, pd.errors.EmptyDataError):
            continue
    raise FileNotFoundError(
        f"Could not load {DATA_PATH}. Export from BigQuery first:\n"
        f"  Get-Content sql\\verify_task_productivity.sql | bq query --use_legacy_sql=false --format=csv > output\\verify_task_productivity.csv"
    )


def safe_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def _compute_duration_minutes(df: pd.DataFrame) -> pd.Series:
    """
    Combined duration: Automated = ceil(task_ended - task_started) to 1 sec, then /60;
    Manual = logic_adjusted_new_step_duration_minutes.
    """
    result = pd.Series(index=df.index, dtype=float)
    logic_adj = pd.to_numeric(df.get("logic_adjusted_new_step_duration_minutes", pd.Series(dtype=float)), errors="coerce")
    is_auto = (df.get("task_mode", pd.Series(dtype=object)) == "Automated").fillna(False).values
    # Manual: use logic_adjusted
    result.loc[~is_auto] = logic_adj.loc[~is_auto]
    # Automated: use task_started_at and task_ended_at, ceil to 1 sec, convert to minutes
    if is_auto.any() and "task_started_at" in df.columns and "task_ended_at" in df.columns:
        started = pd.to_datetime(df.loc[is_auto, "task_started_at"], errors="coerce")
        ended = pd.to_datetime(df.loc[is_auto, "task_ended_at"], errors="coerce")
        diff_sec = (ended - started).dt.total_seconds().clip(lower=0)
        # Round up to nearest 1 second; sub-second (milliseconds) -> at least 1 second
        ceil_sec = np.maximum(np.ceil(diff_sec), 1)
        result.loc[is_auto] = ceil_sec / 60  # 1 second = 1/60 minutes
    return result


def _sample_for_hist(series: pd.Series) -> pd.Series:
    """Subsample if large for faster histogram plotting; preserves distribution shape."""
    if len(series) <= HIST_SAMPLE_MAX:
        return series
    return series.sample(n=HIST_SAMPLE_MAX, random_state=42)


def _safe_filename(label: str) -> str:
    """Sanitize label for use in filename."""
    s = str(label).replace(" | ", "_")
    return "".join(c if c.isalnum() or c in "_-" else "_" for c in s).strip("_") or "unnamed"


def is_bimodal(series: pd.Series, min_n: int = 30) -> bool:
    """
    Detect if distribution is bimodal using two methods:
    1) Bimodality coefficient: BC = (skew² + 1) / (excess_kurtosis + 3). BC > 5/9 suggests bimodality.
    2) KDE peak count: if density has 2+ distinct peaks, likely bimodal.
    """
    s = series.dropna()
    s = s[s > 0]
    if len(s) < min_n:
        return False
    try:
        # Method 1: Bimodality coefficient
        skew = stats.skew(s)
        kurt = stats.kurtosis(s)  # excess kurtosis (0 for normal)
        bc = (skew**2 + 1) / (kurt + 3)
        if bc > 5 / 9:
            return True
        # Method 2: KDE peaks
        kde = stats.gaussian_kde(s)
        x = np.linspace(s.min(), s.max(), 200)
        density = kde(x)
        peaks, _ = find_peaks(density, prominence=density.std() * 0.3)
        return len(peaks) >= 2
    except Exception:
        return False


def _plot_single_histogram(subset: pd.Series, label: str, prefix: str) -> Path:
    """Plot one histogram to its own image file. Labels added for bins between 0-2 min."""
    to_plot = _sample_for_hist(subset)
    fig, ax = plt.subplots(figsize=(8, 5))
    n, bins, patches = ax.hist(to_plot, bins=50, edgecolor="black", alpha=0.7)
    # Add value labels on bars between 0 and 2 minutes
    for i, (count, patch) in enumerate(zip(n, patches)):
        if count <= 0:
            continue
        bin_left = bins[i]
        bin_right = bins[i + 1]
        bin_center = (bin_left + bin_right) / 2
        if 0 <= bin_center <= 2:
            ax.text(
                patch.get_x() + patch.get_width() / 2,
                patch.get_height() + max(n) * 0.01,
                f"{int(count):,}",
                ha="center",
                va="bottom",
                fontsize=8,
            )
    ax.set_title(f"{label}\n(n={len(subset):,})")
    ax.set_xlabel("Duration (minutes)")
    ax.set_ylabel("Count")
    plt.tight_layout()
    fname = _safe_filename((prefix + "_" + label) if prefix else label)
    out = OUTPUT_DIR / f"verify_histogram_{fname}.png"
    plt.savefig(out, dpi=150, bbox_inches="tight")
    plt.close()
    return out


def _plot_manual_first_minute_breakdown(df: pd.DataFrame, duration: pd.Series) -> Path | None:
    """
    Manual tasks only, duration <= 1 min: breakdown by second buckets, emphasizing <=1 sec.
    """
    if "task_mode" not in df.columns:
        return None
    manual = duration[(df["task_mode"] == "Manual")].dropna()
    manual = manual[(manual > 0) & (manual <= 1)]  # 0 < duration <= 1 minute
    if len(manual) == 0:
        return None
    # Convert to seconds for bucketing
    sec = (manual * 60).clip(upper=60)
    # Buckets: <=1, 1-2, 2-5, 5-10, 10-30, 30-60 sec
    bins = [0, 1, 2, 5, 10, 30, 60]
    labels = ["≤1 sec", "1-2 sec", "2-5 sec", "5-10 sec", "10-30 sec", "30-60 sec"]
    bucketed = pd.cut(sec, bins=bins, labels=labels, include_lowest=True)
    counts = bucketed.value_counts().reindex(labels, fill_value=0)
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = ["#e74c3c" if lb == "≤1 sec" else "#3498db" for lb in labels]
    bars = ax.bar(counts.index.astype(str), counts.values, color=colors, edgecolor="black")
    ax.set_title(f"Manual Tasks: Duration ≤1 Minute — Breakdown by Second Buckets\n(n={len(manual):,}, ≤1 sec: {counts.iloc[0]:,} = {100*counts.iloc[0]/len(manual):.1f}%)")
    ax.set_xlabel("Duration (seconds)")
    ax.set_ylabel("Count")
    for bar, v in zip(bars, counts.values):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(counts)*0.01, f"{v:,}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    out = OUTPUT_DIR / "verify_manual_first_minute_breakdown.png"
    plt.savefig(out, dpi=150, bbox_inches="tight")
    plt.close()
    return out


def _plot_topline_counts_by_task_method_warehouse(df: pd.DataFrame) -> Path | None:
    """
    Topline counts bar chart by Task, Task Method, and Warehouse.
    Grouped bars: Task × Warehouse on x-axis, Task Method (Automated/Manual) as hue.
    """
    if "task_mode" not in df.columns or "wh_name" not in df.columns or "task" not in df.columns:
        return None
    counts = (
        df.groupby(["task", "task_mode", "wh_name"])
        .size()
        .reset_index(name="count")
    )
    if counts.empty:
        return None
    counts["task_wh"] = counts["task"].astype(str) + " | " + counts["wh_name"].astype(str)
    fig, ax = plt.subplots(figsize=(10, 6))
    sns.barplot(
        data=counts,
        x="task_wh",
        y="count",
        hue="task_mode",
        palette={"Automated": "#3498db", "Manual": "#2ecc71"},
        edgecolor="black",
        ax=ax,
    )
    ax.set_title("Topline Counts by Task, Task Method, and Warehouse")
    ax.set_xlabel("Task | Warehouse")
    ax.set_ylabel("Count")
    ax.tick_params(axis="x", rotation=45)
    for container in ax.containers:
        ax.bar_label(container, fmt="%d", label_type="edge", fontsize=9)
    plt.legend(title="Task Method")
    plt.tight_layout()
    out = OUTPUT_DIR / "verify_topline_counts_by_task_method_warehouse.png"
    plt.savefig(out, dpi=150, bbox_inches="tight")
    plt.close()
    return out


def _plot_single_kde(subset: pd.Series, label: str, prefix: str) -> Path:
    """Plot one KDE to its own image file."""
    to_plot = _sample_for_hist(subset)
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.kdeplot(data=to_plot, ax=ax, fill=True, alpha=0.5)
    ax.set_title(f"{label}\n(n={len(subset):,})")
    ax.set_xlabel("Duration (minutes)")
    ax.set_ylabel("Density")
    ax.set_xlim(left=0)
    plt.tight_layout()
    fname = _safe_filename((prefix + "_" + label) if prefix else label)
    out = OUTPUT_DIR / f"verify_kde_{fname}.png"
    plt.savefig(out, dpi=150, bbox_inches="tight")
    plt.close()
    return out


def generate_all_plots(df: pd.DataFrame, duration: pd.Series) -> list[Path]:
    """Generate individual histogram and KDE images for each cut. Returns list of paths."""
    images = []
    # Topline counts by Task, Task Method, Warehouse
    p = _plot_topline_counts_by_task_method_warehouse(df)
    if p:
        images.append(p)
    # Overall
    d = duration.dropna()
    d = d[d > 0]
    if len(d) > 0:
        images.append(_plot_single_histogram(d, "Overall", ""))
        images.append(_plot_single_kde(d, "Overall", ""))
    # By task
    for task in df["task"].dropna().unique():
        subset = duration[df["task"] == task].dropna()
        subset = subset[subset > 0]
        if len(subset) > 0:
            images.append(_plot_single_histogram(subset, str(task), "task"))
            images.append(_plot_single_kde(subset, str(task), "task"))
    # By task_mode (Automated vs Manual)
    if "task_mode" in df.columns:
        for mode in df["task_mode"].dropna().unique():
            subset = duration[df["task_mode"] == mode].dropna()
            subset = subset[subset > 0]
            if len(subset) > 0:
                images.append(_plot_single_histogram(subset, str(mode), "task_mode"))
                images.append(_plot_single_kde(subset, str(mode), "task_mode"))
    # By task_mode × task (Automated/Manual - 10 Pending Verify, 06 Preverify)
    if "task_mode" in df.columns:
        for mode in ["Automated", "Manual"]:
            for task in ["10 Pending Verify", "06 Preverify"]:
                mask = (df["task_mode"] == mode) & (df["task"] == task)
                subset = duration[mask].dropna()
                subset = subset[subset > 0]
                if len(subset) > 0:
                    label = f"{mode} - {task}"
                    images.append(_plot_single_histogram(subset, label, "mode_task"))
                    images.append(_plot_single_kde(subset, label, "mode_task"))
                elif mask.sum() > 0:
                    print(f"  Skipped {mode} - {task}: {mask.sum():,} rows but all have duration<=0 (automated tasks often have 0 duration)")
    # By wh_name
    if "wh_name" in df.columns:
        for wh in df["wh_name"].dropna().unique():
            subset = duration[df["wh_name"] == wh].dropna()
            subset = subset[subset > 0]
            if len(subset) > 0:
                images.append(_plot_single_histogram(subset, str(wh), "wh"))
                images.append(_plot_single_kde(subset, str(wh), "wh"))
    # By task × wh_name (top 8)
    if "wh_name" in df.columns:
        df = df.copy()
        df["task_wh"] = df["task"].astype(str) + " | " + df["wh_name"].astype(str)
        for combo in df.groupby("task_wh").size().nlargest(8).index:
            subset = duration[df["task_wh"] == combo].dropna()
            subset = subset[subset > 0]
            if len(subset) > 0:
                images.append(_plot_single_histogram(subset, str(combo), "task_wh"))
                images.append(_plot_single_kde(subset, str(combo), "task_wh"))
    # Manual only: first minute breakdown (≤1 min by second buckets, emphasize ≤1 sec)
    p = _plot_manual_first_minute_breakdown(df, duration)
    if p:
        images.append(p)
    return images


def build_summary_cuts(df: pd.DataFrame, duration: pd.Series) -> pd.DataFrame:
    """Summary stats by cut dimensions."""
    rows = []
    # Overall
    d = duration.dropna()
    d = d[d > 0]
    rows.append({
        "cut": "Overall",
        "dimension": "—",
        "value": "—",
        "n": len(d),
        "mean_min": round(d.mean(), 2),
        "median_min": round(d.median(), 2),
        "p25": round(d.quantile(0.25), 2),
        "p75": round(d.quantile(0.75), 2),
        "p90": round(d.quantile(0.90), 2),
        "bimodal": is_bimodal(d),
    })
    # By task
    for task in df["task"].dropna().unique():
        d = duration[df["task"] == task].dropna()
        d = d[d > 0]
        if len(d) > 0:
            rows.append({
                "cut": "task",
                "dimension": "task",
                "value": str(task),
                "n": len(d),
                "mean_min": round(d.mean(), 2),
                "median_min": round(d.median(), 2),
                "p25": round(d.quantile(0.25), 2),
                "p75": round(d.quantile(0.75), 2),
                "p90": round(d.quantile(0.90), 2),
                "bimodal": is_bimodal(d),
            })
    # By task_mode (Automated vs Manual: beyondrx email = Automated)
    if "task_mode" in df.columns:
        for mode in df["task_mode"].dropna().unique():
            d = duration[df["task_mode"] == mode].dropna()
            d = d[d > 0]
            if len(d) > 0:
                rows.append({
                    "cut": "task_mode",
                    "dimension": "task_mode",
                    "value": str(mode),
                    "n": len(d),
                    "mean_min": round(d.mean(), 2),
                    "median_min": round(d.median(), 2),
                    "p25": round(d.quantile(0.25), 2),
                    "p75": round(d.quantile(0.75), 2),
                    "p90": round(d.quantile(0.90), 2),
                    "bimodal": is_bimodal(d),
                })
    # By wh_name
    if "wh_name" in df.columns:
        for wh in df["wh_name"].dropna().unique():
            d = duration[df["wh_name"] == wh].dropna()
            d = d[d > 0]
            if len(d) > 0:
                rows.append({
                    "cut": "wh_name",
                    "dimension": "wh_name",
                    "value": str(wh),
                    "n": len(d),
                    "mean_min": round(d.mean(), 2),
                    "median_min": round(d.median(), 2),
                    "p25": round(d.quantile(0.25), 2),
                    "p75": round(d.quantile(0.75), 2),
                    "p90": round(d.quantile(0.90), 2),
                    "bimodal": is_bimodal(d),
                })
    # By task_mode × task (Automated/Manual - 10 Pending Verify, 06 Preverify)
    if "task_mode" in df.columns:
        for mode in ["Automated", "Manual"]:
            for task in ["10 Pending Verify", "06 Preverify"]:
                mask = (df["task_mode"] == mode) & (df["task"] == task)
                d = duration[mask].dropna()
                d = d[d > 0]
                if len(d) > 0:
                    rows.append({
                        "cut": "task_mode × task",
                        "dimension": "mode_task",
                        "value": f"{mode} - {task}",
                        "n": len(d),
                        "mean_min": round(d.mean(), 2),
                        "median_min": round(d.median(), 2),
                        "p25": round(d.quantile(0.25), 2),
                        "p75": round(d.quantile(0.75), 2),
                        "p90": round(d.quantile(0.90), 2),
                        "bimodal": is_bimodal(d),
                    })
                elif mask.sum() > 0:
                    rows.append({
                        "cut": "task_mode × task",
                        "dimension": "mode_task",
                        "value": f"{mode} - {task}",
                        "n": mask.sum(),
                        "mean_min": "—",
                        "median_min": "—",
                        "p25": "—",
                        "p75": "—",
                        "p90": "—",
                        "bimodal": False,
                    })
    # By task × wh_name (top combos)
    if "wh_name" in df.columns:
        df["task_wh"] = df["task"].astype(str) + " | " + df["wh_name"].astype(str)
        for combo in df.groupby("task_wh").size().nlargest(6).index:
            d = duration[df["task_wh"] == combo].dropna()
            d = d[d > 0]
            if len(d) > 0:
                rows.append({
                    "cut": "task × wh_name",
                    "dimension": "task_wh",
                    "value": str(combo),
                    "n": len(d),
                    "mean_min": round(d.mean(), 2),
                    "median_min": round(d.median(), 2),
                    "p25": round(d.quantile(0.25), 2),
                    "p75": round(d.quantile(0.75), 2),
                    "p90": round(d.quantile(0.90), 2),
                    "bimodal": is_bimodal(d),
                })
    return pd.DataFrame(rows)


# Bin edges (seconds) and labels for Task × Task Method tables
DURATION_BIN_EDGES_SEC = [0, 1, 10, 20, 30, 60, 120, 180, 240, 300, np.inf]
DURATION_BIN_LABELS = [
    "Less than 1 sec",
    "1-10 sec",
    "11-20 sec",
    "21-30 sec",
    "31-60 sec",
    "1-2 min",
    "2-3 min",
    "3-4 min",
    "4-5 min",
    "5+ min",
]


def build_task_mode_bin_tables(df: pd.DataFrame, duration: pd.Series) -> pd.DataFrame:
    """
    For each Task × Task Method combination, count observations in duration bins:
    Less than 1 sec, 1-10 sec, 11-20 sec, 21-30 sec, 31-60 sec, 1-2 min, 2-3 min, 3-4 min, 4-5 min, 5+ min.
    """
    if "task_mode" not in df.columns or "task" not in df.columns:
        return pd.DataFrame()
    sec = (duration.dropna() * 60).replace([np.inf, -np.inf], np.nan)
    sec = sec[sec > 0]
    df_sub = df.loc[sec.index].copy()
    df_sub["duration_sec"] = sec
    df_sub = df_sub[df_sub["duration_sec"].notna()]
    df_sub["bin"] = pd.cut(
        df_sub["duration_sec"],
        bins=DURATION_BIN_EDGES_SEC,
        labels=DURATION_BIN_LABELS,
        include_lowest=True,
    )
    rows = []
    for (task, mode), grp in df_sub.groupby(["task", "task_mode"]):
        total = len(grp)
        counts = grp["bin"].value_counts().reindex(DURATION_BIN_LABELS, fill_value=0)
        for label in DURATION_BIN_LABELS:
            count = int(counts.get(label, 0))
            pct = 100 * count / total if total > 0 else 0
            rows.append({
                "task": task,
                "task_mode": mode,
                "bin": label,
                "count": count,
                "pct": round(pct, 1),
                "total": total,
            })
    return pd.DataFrame(rows)


def build_doc(
    summary_df: pd.DataFrame,
    image_paths: list[Path],
    bin_tables_df: pd.DataFrame | None = None,
) -> Path:
    """Build Word document with cuts and visuals."""
    doc = Document()
    doc.add_heading("Task Productivity: Verify Tasks — Duration Distribution", 0)
    doc.add_paragraph(
        "Histogram analysis of duration (minutes). Automated tasks: ceil(task_ended - task_started) to 1 sec, then /60. "
        "Manual tasks: logic_adjusted_new_step_duration_minutes. "
        "Cuts: task, wh_name, task_mode (Automated=BeyondRx email, Manual=otherwise), task_mode × task."
    )
    doc.add_paragraph()

    # Duration bins by Task × Task Method
    if bin_tables_df is not None and not bin_tables_df.empty:
        doc.add_heading("Duration Bins by Task × Task Method", level=1)
        doc.add_paragraph(
            "Bins: Less than 1 sec, 1-10 sec, 11-20 sec, 21-30 sec, 31-60 sec, 1-2 min, 2-3 min, 3-4 min, 4-5 min, 5+ min."
        )
        for (task, mode), grp in bin_tables_df.groupby(["task", "task_mode"]):
            doc.add_heading(f"{mode} — {task}", level=2)
            tbl = grp.set_index("bin").reindex(DURATION_BIN_LABELS).reset_index()
            tbl = tbl[["bin", "count", "pct"]].fillna(0)
            tbl.columns = ["Duration bin", "Count", "%"]
            t = doc.add_table(rows=len(tbl) + 1, cols=3, style="Table Grid")
            t.rows[0].cells[0].text = "Duration bin"
            t.rows[0].cells[1].text = "Count"
            t.rows[0].cells[2].text = "%"
            for idx, (_, row) in enumerate(tbl.iterrows()):
                r = t.rows[idx + 1]
                r.cells[0].text = str(row["Duration bin"])
                r.cells[1].text = str(int(row["Count"]))
                r.cells[2].text = str(round(row["%"], 1)) if pd.notna(row["%"]) else "0"
            doc.add_paragraph()

    doc.add_heading("Summary Statistics by Cut", level=1)
    doc.add_paragraph(
        "n = count of observations; mean_min, median_min, p25, p75, p90 in minutes. "
        "bimodal = True if distribution appears bimodal. "
        "Note: Automated tasks (beyondrx email) often have duration=0 and are excluded from duration stats; they appear with n only."
    )
    if not summary_df.empty:
        t = doc.add_table(rows=len(summary_df) + 1, cols=len(summary_df.columns), style="Table Grid")
        for j, col in enumerate(summary_df.columns):
            t.rows[0].cells[j].text = str(col)
        for i, row in summary_df.iterrows():
            for j, val in enumerate(row):
                t.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()

    doc.add_heading("Visualizations", level=1)
    for p in image_paths:
        if p and p.exists():
            doc.add_heading(p.stem.replace("_", " ").title(), level=2)
            doc.add_picture(str(p), width=Inches(5.5))
            doc.add_paragraph()

    out = OUTPUT_DIR / "verify_task_productivity_report.docx"
    doc.save(str(out))
    return out


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    t0 = time.perf_counter()
    print("Loading data...")
    df = load_data()
    if DURATION_COL not in df.columns:
        raise ValueError(f"Column '{DURATION_COL}' not found. Columns: {list(df.columns)}")
    duration = pd.to_numeric(df[DURATION_COL], errors="coerce")
    print(f"  Rows: {len(df):,}, valid duration: {duration.notna().sum():,} ({time.perf_counter()-t0:.1f}s)")
    if "task_mode" in df.columns:
        mode_counts = df["task_mode"].value_counts()
        print(f"  task_mode: {dict(mode_counts)}")

    t1 = time.perf_counter()
    print("Building summary cuts...")
    summary_df = build_summary_cuts(df, duration)
    summary_df.to_csv(OUTPUT_DIR / "verify_summary_cuts.csv", index=False)
    print(f"  Saved: {OUTPUT_DIR / 'verify_summary_cuts.csv'} ({time.perf_counter()-t1:.1f}s)")

    t1b = time.perf_counter()
    print("Building Task × Task Method bin tables...")
    bin_tables_df = build_task_mode_bin_tables(df, duration)
    if not bin_tables_df.empty:
        bin_tables_df.to_csv(OUTPUT_DIR / "verify_duration_bins_by_task_mode.csv", index=False)
        print(f"  Saved: {OUTPUT_DIR / 'verify_duration_bins_by_task_mode.csv'} ({time.perf_counter()-t1b:.1f}s)")

    t2 = time.perf_counter()
    print("Generating histograms and KDE plots (one image per cut)...")
    images = generate_all_plots(df, duration)
    images = sorted(
        [p for p in images if p.exists()],
        key=lambda p: (0 if "topline" in p.name else 1, not p.name.startswith("verify_histogram"), p.name),
    )
    for p in images:
        print(f"  Saved: {p}")
    print(f"  Plots done ({time.perf_counter()-t2:.1f}s)")

    t3 = time.perf_counter()
    print("Building Word document...")
    out_doc = build_doc(summary_df, images, bin_tables_df)
    print(f"  Saved: {out_doc} ({time.perf_counter()-t3:.1f}s)")
    print(f"Done. Total: {time.perf_counter()-t0:.1f}s")


if __name__ == "__main__":
    main()
