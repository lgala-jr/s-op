"""
Compile all dispense method analysis outputs into a Word document.
Run after: base data export, feature selection, and SQL summary/distribution queries.
"""

import argparse
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path("output")
FEATURE_SELECTION_DIR = OUTPUT_DIR / "feature_selection"
REPORT_PATH = OUTPUT_DIR / "Dispense_Method_Analysis_Report.docx"

# Expected output paths (from SQL exports)
EXPECTED_CSVS = {
    "summary_stats": OUTPUT_DIR / "summary_stats.csv",
    "summary_by_dimensions": OUTPUT_DIR / "summary_by_dimensions.csv",
    "distribution_by_product_type": OUTPUT_DIR / "distribution_by_product_type.csv",
    "distribution_by_wh": OUTPUT_DIR / "distribution_by_wh.csv",
    "chi_square_scores": FEATURE_SELECTION_DIR / "chi_square_scores.csv",
    "mutual_info_scores": FEATURE_SELECTION_DIR / "mutual_info_scores.csv",
}

EXPECTED_IMAGES = [
    ("Dispense Overview", FEATURE_SELECTION_DIR / "00_dispense_overview.png"),
    ("Feature Distributions", FEATURE_SELECTION_DIR / "01_feature_distributions.png"),
    ("Association Strength", FEATURE_SELECTION_DIR / "03_association_strength.png"),
]


def read_csv_safe(path: Path) -> pd.DataFrame | None:
    """Read CSV with encoding fallback. Returns None if empty or invalid."""
    for enc in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
        try:
            df = pd.read_csv(path, encoding=enc)
            if df.empty or len(df.columns) == 0:
                return None
            return df
        except (UnicodeDecodeError, pd.errors.EmptyDataError, pd.errors.ParserError):
            continue
    return None


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def df_to_table(doc: Document, df: pd.DataFrame, style: str = "Table Grid") -> None:
    """Insert a pandas DataFrame as a Word table."""
    df = df.fillna("")
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style=style)
    # Header row
    for j, col in enumerate(df.columns):
        table.rows[0].cells[j].text = str(col)
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    # Data rows
    for i, row in df.iterrows():
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str | None = None, width: float = 5.5) -> None:
    """Insert an image with optional caption."""
    if not path.exists():
        doc.add_paragraph(f"[Image not found: {path}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    doc.add_paragraph()


def add_feature_vs_dispense_images(doc: Document) -> None:
    """Add all 02_<feature>_vs_dispense.png images."""
    for f in FEATURE_SELECTION_DIR.glob("02_*_vs_dispense.png"):
        feat_name = f.stem.replace("02_", "").replace("_vs_dispense", "").replace("_", " ").title()
        add_heading(doc, f"{feat_name} vs Dispense Method", level=3)
        add_image(doc, f, caption=f"{feat_name} distribution by dispense_method", width=5.0)


def main():
    parser = argparse.ArgumentParser(description="Generate Word report from analysis outputs")
    parser.add_argument(
        "-o", "--output",
        default=str(REPORT_PATH),
        help="Output Word document path",
    )
    args = parser.parse_args()

    doc = Document()

    # Title
    doc.add_heading("Dispense Method Analysis Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    # --- 1. Feature Selection ---
    add_heading(doc, "1. Feature Selection", level=1)

    # Selected features
    selected_path = FEATURE_SELECTION_DIR / "selected_features.txt"
    if selected_path.exists():
        add_heading(doc, "Selected Features", level=2)
        with open(selected_path) as f:
            lines = [l.strip() for l in f if l.strip() and not l.startswith("#")]
        doc.add_paragraph(", ".join(lines) if lines else "(none)")
        doc.add_paragraph()

    # Chi-square scores with explanation
    add_heading(doc, "Chi-Square Scores", level=2)
    doc.add_paragraph(
        "The chi-square test measures the strength of association between each candidate feature and dispense_method. "
        "It tests the null hypothesis that the feature and target are independent. A higher chi-square value indicates "
        "stronger association; a low p-value (typically < 0.05) means we reject independence and conclude the feature "
        "is useful for predicting dispense method. Chi-square is well-suited for categorical variables and helps identify "
        "which product attributes (e.g., product_type, unit_of_measure) are most predictive of how a product will be dispensed."
    )
    csv_path = FEATURE_SELECTION_DIR / "chi_square_scores.csv"
    if csv_path.exists():
        df = read_csv_safe(csv_path)
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Chi-square scores not found — run feature_selection.py first]")
        doc.add_paragraph()

    # Mutual information scores with explanation
    add_heading(doc, "Mutual Information Scores", level=2)
    doc.add_paragraph(
        "Mutual information (MI) measures how much knowing the value of a feature reduces uncertainty about dispense_method. "
        "Unlike chi-square, MI captures non-linear relationships and does not assume a specific distribution. A higher MI "
        "score means the feature provides more information about the target. MI is useful for ranking features when building "
        "prediction models because it identifies which attributes (e.g., level_3, product_type) best reduce uncertainty about "
        "whether a product will be dispensed via automation, manual, or mixture. Features with high MI are strong candidates "
        "for the prediction model."
    )
    csv_path = FEATURE_SELECTION_DIR / "mutual_info_scores.csv"
    if csv_path.exists():
        df = read_csv_safe(csv_path)
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Mutual information scores not found — run feature_selection.py first]")
        doc.add_paragraph()

    # Feature selection images
    add_heading(doc, "Feature Selection Visualizations", level=2)
    for caption, img_path in EXPECTED_IMAGES:
        add_heading(doc, caption, level=3)
        add_image(doc, img_path, width=5.5)
    add_feature_vs_dispense_images(doc)

    # --- 2. Summary Statistics ---
    add_heading(doc, "2. Summary Statistics", level=1)

    if EXPECTED_CSVS["summary_stats"].exists():
        add_heading(doc, "Overall Dispense Method Summary", level=2)
        df = read_csv_safe(EXPECTED_CSVS["summary_stats"])
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Run 02_summary_stats.sql and export to output/summary_stats.csv]")
        doc.add_paragraph()

    if EXPECTED_CSVS["summary_by_dimensions"].exists():
        add_heading(doc, "Summary by Product Type and Dispense Method", level=2)
        df = read_csv_safe(EXPECTED_CSVS["summary_by_dimensions"])
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Run 02b_summary_by_dimensions.sql and export to output/summary_by_dimensions.csv]")
        doc.add_paragraph()

    # --- 3. Distributions ---
    add_heading(doc, "3. Distribution Analysis", level=1)

    if EXPECTED_CSVS["distribution_by_product_type"].exists():
        add_heading(doc, "Distribution by Product Type", level=2)
        df = read_csv_safe(EXPECTED_CSVS["distribution_by_product_type"])
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Run 03_distributions.sql and export to output/distribution_by_product_type.csv]")
        doc.add_paragraph()

    if EXPECTED_CSVS["distribution_by_wh"].exists():
        add_heading(doc, "Distribution by Warehouse (wh_name)", level=2)
        df = read_csv_safe(EXPECTED_CSVS["distribution_by_wh"])
        if df is not None:
            df_to_table(doc, df)
        else:
            doc.add_paragraph("[File empty or invalid]")
    else:
        doc.add_paragraph("[Run 03b_distribution_by_wh.sql and export to output/distribution_by_wh.csv]")
        doc.add_paragraph()

    # --- 4. Dispense Ratio Lookup (Forecast Allocation) ---
    ratio_path = OUTPUT_DIR / "dispense_ratio_lookup.csv"
    if ratio_path.exists():
        add_heading(doc, "4. Dispense Ratio Lookup (Apply to Forecast Volume)", level=1)
        doc.add_paragraph(
            "Historical ratios from prior 6 months. Use to allocate forecast volume across dispense methods. "
            "Level 1 = product_type × wh_name (preferred); fall back to level 2 (product_type), 3 (wh_name), or 4 (overall) if no match."
        )
        doc.add_paragraph("Formula: estimated_dispense_volume = forecast_volume × (pct_of_volume / 100)")
        df = read_csv_safe(ratio_path)
        if df is not None:
            df_to_table(doc, df)
        doc.add_paragraph("See RATIO_APPLICATION.md for full usage.")
    else:
        doc.add_paragraph("[Run build_dispense_ratios.py to generate dispense_ratio_lookup.csv]")
        doc.add_paragraph()

    # --- 5. Next Steps ---
    add_heading(doc, "5. Next Steps", level=1)
    doc.add_paragraph(
        "• Apply dispense_ratio_lookup.csv to forecast volume (product_type × wh_name)\n"
        "• Refresh ratios monthly/quarterly by re-running the analysis pipeline\n"
        "• Use level 1 ratios first; fall back to level 2/3/4 when forecast lacks dimensions"
    )

    # Save
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Report saved to {out_path}")


if __name__ == "__main__":
    main()
